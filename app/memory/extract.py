# memory/extract.py
from __future__ import annotations

import csv
import hashlib
import io
import json
import logging
import mimetypes
import os
import re
import shutil
import subprocess
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ---------- Optional deps (import lazily/optionally) ----------
try:
    import magic  # python-magic (libmagic)
except Exception:  # pragma: no cover
    magic = None

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    import docx  # python-docx
except Exception:  # pragma: no cover
    docx = None

try:
    import openpyxl  # XLSX/XLSM (stream/read_only supported)
except Exception:  # pragma: no cover
    openpyxl = None

try:
    import xlrd  # legacy .xls support (requires <2.0)
except Exception:  # pragma: no cover
    xlrd = None

try:
    import chardet  # optional charset detection
except Exception:  # pragma: no cover
    chardet = None


# ---------- Config ----------
DEFAULT_TXT_ENCODING = "utf-8"
SUBPROCESS_TIMEOUT = 60  # seconds
MAX_SOFT_BYTES_TXT = None  # set int to cap giant .txt reads; None = full read
MAX_SAMPLE_ROWS_PER_SHEET = 50
MAX_SAMPLE_ROWS_CSV = 100
READ_SAMPLE_BYTES_FOR_SNIFF = 8192


# ---------- Types ----------
@dataclass(frozen=True)
class ExtractionError(Exception):
    path: str
    reason: str

    def __str__(self) -> str:  # pragma: no cover
        return f"ExtractionError({self.path}): {self.reason}"


@dataclass
class ExtractionStats:
    size_bytes: int
    num_chars: int
    num_words: int
    num_lines: int


@dataclass
class ExtractionResult:
    path: str
    mime: str
    ext: str
    encoding: Optional[str]
    text: str  # sanitized plain text (may be sample-limited for tabulars)
    stats: ExtractionStats
    structure: Dict[str, Any]

    def to_dict(self) -> Dict[str, Any]:
        d = asdict(self)
        d["stats"] = asdict(self.stats)
        return d


# ---------- Supported families ----------
_EXT_TO_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
    ".xls": "application/vnd.ms-excel",
    ".csv": "text/csv",
    ".tsv": "text/tab-separated-values",
    ".txt": "text/plain",
    ".json": "application/json",
}
CSV_EXTS = {".csv", ".tsv"}
EXCEL_EXTS = {".xlsx", ".xlsm", ".xls"}
DOC_EXTS = {".doc", ".docx"}
TXT_EXTS = {".txt"}
PDF_EXTS = {".pdf"}
JSON_EXTS = {".json"}

SUPPORTED_EXTS = CSV_EXTS | EXCEL_EXTS | DOC_EXTS | TXT_EXTS | PDF_EXTS | JSON_EXTS


# ---------- Utilities ----------
def _run(cmd: List[str], *, timeout: int = SUBPROCESS_TIMEOUT) -> str:
    """Run a command and return stdout as UTF-8 (lossy)."""
    logger.debug("Running command: %s", " ".join(cmd))
    out = subprocess.check_output(cmd, stderr=subprocess.STDOUT, timeout=timeout)
    return out.decode("utf-8", "ignore")


def _detect_encoding(path: str, default: str = DEFAULT_TXT_ENCODING) -> str:
    if not chardet:
        return default
    try:
        with open(path, "rb") as fh:
            sample = fh.read(READ_SAMPLE_BYTES_FOR_SNIFF)
        guess = chardet.detect(sample)
        enc = guess.get("encoding") or default
        return enc
    except Exception:
        return default


def detect_mime(path: str) -> str:
    """Best-effort MIME detection using libmagic, then curated map, then mimetypes."""
    if magic:
        try:
            with open(path, "rb") as fh:
                return magic.from_buffer(fh.read(4096), mime=True)
        except Exception:
            pass
    ext = Path(path).suffix.lower()
    if ext in _EXT_TO_MIME:
        return _EXT_TO_MIME[ext]
    guess, _ = mimetypes.guess_type(path)
    return guess or "application/octet-stream"


_CTRL_RE = re.compile(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]")


def _sanitize_text(s: Optional[str], *, keep_newlines: bool = True) -> str:
    """Normalize newlines, strip control chars (keeps \\t and \\n)."""
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = _CTRL_RE.sub("", s)
    if not keep_newlines:
        s = s.replace("\n", " ")
    return s


_WORD_RE = re.compile(r"\w+", re.UNICODE)


def _basic_stats(text: str, size_bytes: int) -> ExtractionStats:
    return ExtractionStats(
        size_bytes=size_bytes,
        num_chars=len(text),
        num_words=len(_WORD_RE.findall(text)),
        num_lines=(text.count("\n") + (1 if text and not text.endswith("\n") else 0)),
    )


def _require_ext(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ExtractionError(path, f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTS)}")
    return ext


def _sha256_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


# ---------- Readers ----------
def _read_txt(path: str, encoding: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
    enc = encoding or _detect_encoding(path)
    size = os.path.getsize(path)
    if MAX_SOFT_BYTES_TXT and size > MAX_SOFT_BYTES_TXT:
        logger.warning("Truncating text read for %s to %d bytes", path, MAX_SOFT_BYTES_TXT)
        with open(path, "rb") as fh:
            data = fh.read(MAX_SOFT_BYTES_TXT)
        text = data.decode(enc, "ignore")
    else:
        with open(path, "r", encoding=enc, errors="ignore") as fh:
            text = fh.read()
    meta = {"kind": "text", "encoding": enc}
    return text, meta


def _read_pdf_pymupdf(path: str) -> Tuple[str, Dict[str, Any]]:
    if not fitz:
        raise ExtractionError(path, "PDF support requires PyMuPDF (fitz)")
    doc = fitz.open(path)
    try:
        pages_meta: List[Dict[str, Any]] = []
        all_text_parts: List[str] = []
        total_images = total_links = total_annots = 0

        for i, page in enumerate(doc):
            p_text = page.get_text("text") or ""
            img_count = len(page.get_images(full=True))
            link_count = len(page.get_links() or [])
            ann_count = 0
            try:
                ann_iter = page.annots()
                if ann_iter:
                    for _ in ann_iter:
                        ann_count += 1
            except Exception:
                pass

            total_images += img_count
            total_links += link_count
            total_annots += ann_count
            pages_meta.append(
                {
                    "index": i,
                    "chars": len(p_text),
                    "words": len(_WORD_RE.findall(p_text)),
                    "images": img_count,
                    "links": link_count,
                    "annotations": ann_count,
                }
            )
            all_text_parts.append(p_text)

        outline: List[Dict[str, Any]] = []
        try:
            toc = doc.get_toc(simple=True) or []  # [level, title, page]
            for level, title, page_no in toc:
                outline.append({"level": level, "title": title, "page": int(page_no) - 1})
        except Exception:
            pass

        meta = {
            "kind": "pdf",
            "page_count": len(doc),
            "pages": pages_meta,
            "outline": outline,
            "total_images": total_images,
            "total_links": total_links,
            "total_annotations": total_annots,
            "engine": "pymupdf",
        }
        return "\n".join(all_text_parts), meta
    finally:
        doc.close()


def _read_pdf_pdftotext(path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Fallback using poppler-utils 'pdftotext' if PyMuPDF is unavailable.
    """
    if not shutil.which("pdftotext"):
        raise ExtractionError(path, "No PDF engine available (install PyMuPDF or pdftotext).")
    out = _run(["pdftotext", "-layout", path, "-"])  # preserve layout best-effort
    meta = {"kind": "pdf", "engine": "pdftotext"}
    return out, meta


def _read_pdf(path: str) -> Tuple[str, Dict[str, Any]]:
    try:
        return _read_pdf_pymupdf(path)
    except ExtractionError as e:
        # try fallback only if engine missing
        if "PyMuPDF" in e.reason:
            return _read_pdf_pdftotext(path)
        raise


def _read_docx(path: str) -> Tuple[str, Dict[str, Any]]:
    if not docx:
        raise ExtractionError(path, "DOCX support requires python-docx")
    d = docx.Document(path)

    parts: List[str] = []
    headings: List[Dict[str, Any]] = []
    for p in d.paragraphs:
        txt = p.text or ""
        if txt:
            parts.append(txt)
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            style_name = ""
        if style_name.lower().startswith("heading"):
            m = re.search(r"(\d+)$", style_name)
            level = int(m.group(1)) if m else None
            headings.append({"title": txt, "level": level, "style": style_name})

    tables_meta: List[Dict[str, Any]] = []
    for ti, t in enumerate(d.tables):
        nrows = len(t.rows)
        ncols = len(t.columns) if getattr(t, "columns", None) else (len(t.rows[0].cells) if nrows else 0)
        header: List[str] = []
        if nrows:
            try:
                header = [c.text or "" for c in t.rows[0].cells]
            except Exception:
                header = []
        tables_meta.append({"index": ti, "rows": nrows, "cols": ncols, "header": header})
        for row in t.rows[:MAX_SAMPLE_ROWS_PER_SHEET]:
            parts.append("\t".join(c.text or "" for c in row.cells))

    meta = {
        "kind": "docx",
        "paragraphs": len(d.paragraphs),
        "headings": headings,
        "tables_count": len(d.tables),
        "tables": tables_meta,
    }
    return "\n".join(parts), meta


def _read_doc(path: str) -> Tuple[str, Dict[str, Any]]:
    """
    Read legacy .doc using progressive fallbacks:
      1) antiword
      2) catdoc
      3) LibreOffice headless conversion to TXT
    """
    if shutil.which("antiword"):
        try:
            text = _run(["antiword", path])
            return text, {"kind": "doc", "reader": "antiword"}
        except Exception as e:
            logger.info("antiword failed on %s: %s", path, e)

    if shutil.which("catdoc"):
        try:
            text = _run(["catdoc", "-w", path])
            return text, {"kind": "doc", "reader": "catdoc"}
        except Exception as e:
            logger.info("catdoc failed on %s: %s", path, e)

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        tmp_dir = Path(path).parent
        out_txt = tmp_dir / (Path(path).stem + ".txt")
        try:
            _run([soffice, "--headless", "--convert-to", "txt:Text", "--outdir", str(tmp_dir), path])
            if out_txt.exists():
                try:
                    text, _ = _read_txt(str(out_txt))
                    return text, {"kind": "doc", "reader": "libreoffice"}
                finally:
                    try:
                        out_txt.unlink(missing_ok=True)
                    except Exception:
                        pass
        except Exception as e:
            logger.info("LibreOffice conversion failed on %s: %s", path, e)

    raise ExtractionError(path, "Cannot read .doc; install antiword/catdoc or LibreOffice")


def _read_xlsx_like(path: str) -> Tuple[str, Dict[str, Any]]:
    """
    XLSX/XLSM via openpyxl (streaming read); XLS via xlrd.
    Returns a TSV-ish text sample plus full structural counts.
    """
    ext = Path(path).suffix.lower()
    if ext in {".xlsx", ".xlsm"}:
        if not openpyxl:
            raise ExtractionError(path, "XLSX/XLSM support requires openpyxl")
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        try:
            sheets_meta: List[Dict[str, Any]] = []
            text_parts: List[str] = []
            for ws in wb.worksheets:
                rows = 0
                cols_max = 0
                header: List[str] = []
                sample_rows = 0
                text_parts.append(f"# Sheet: {ws.title}")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    values = ["" if c is None else str(c) for c in row]
                    rows += 1
                    cols_max = max(cols_max, len(values))
                    if i == 0:
                        header = values
                    if sample_rows < MAX_SAMPLE_ROWS_PER_SHEET:
                        text_parts.append("\t".join(values))
                        sample_rows += 1
                sheets_meta.append({"title": ws.title, "rows": rows, "cols": cols_max, "header": header})
            meta = {"kind": "excel", "engine": "openpyxl", "sheet_count": len(sheets_meta), "sheets": sheets_meta}
            return "\n".join(text_parts), meta
        finally:
            wb.close()

    if ext == ".xls":
        if not xlrd:
            raise ExtractionError(path, "XLS support requires xlrd<2.0 (legacy .xls)")
        book = xlrd.open_workbook(path)
        sheets_meta: List[Dict[str, Any]] = []
        text_parts: List[str] = []
        for si in range(book.nsheets):
            sh = book.sheet_by_index(si)
            rows = sh.nrows
            cols = sh.ncols
            header = [str(sh.cell_value(0, c)) for c in range(cols)] if rows else []
            sheets_meta.append({"title": sh.name, "rows": rows, "cols": cols, "header": header})
            text_parts.append(f"# Sheet: {sh.name}")
            for r in range(min(rows, MAX_SAMPLE_ROWS_PER_SHEET)):
                values = [str(sh.cell_value(r, c)) for c in range(cols)]
                text_parts.append("\t".join(values))
        meta = {"kind": "excel", "engine": "xlrd", "sheet_count": len(sheets_meta), "sheets": sheets_meta}
        return "\n".join(text_parts), meta

    raise ExtractionError(path, f"Unsupported Excel extension '{ext}'")


def _read_csv_tsv_stream(path: str, delimiter: Optional[str] = None) -> Tuple[str, Dict[str, Any]]:
    enc = _detect_encoding(path)
    row_count = 0
    cols_max = 0
    header: List[str] = []
    has_header = False
    sample_lines = 0
    sample_buffer = io.StringIO()

    with open(path, "r", encoding=enc, errors="ignore", newline="") as fh:
        head = fh.read(READ_SAMPLE_BYTES_FOR_SNIFF)
        fh.seek(0)
        if delimiter is None:
            try:
                dialect = csv.Sniffer().sniff(head)
                delimiter = dialect.delimiter
            except Exception:
                delimiter = "\t" if path.lower().endswith(".tsv") else ","
        try:
            has_header = csv.Sniffer().has_header(head)
        except Exception:
            has_header = False

        reader = csv.reader(fh, delimiter=delimiter)
        for i, row in enumerate(reader):
            row_count += 1
            cols_max = max(cols_max, len(row))
            if i == 0:
                header = row[:]
            if sample_lines < MAX_SAMPLE_ROWS_CSV:
                sample_buffer.write(delimiter.join(row))
                sample_buffer.write("\n")
                sample_lines += 1

    meta = {
        "kind": "csv",
        "encoding": enc,
        "delimiter": delimiter,
        "has_header": bool(has_header),
        "rows": row_count,
        "cols": cols_max,
        "header": header,
    }
    return sample_buffer.getvalue(), meta


def _read_json(path: str) -> Tuple[str, Dict[str, Any]]:
    enc = _detect_encoding(path)
    with open(path, "r", encoding=enc, errors="ignore") as fh:
        data = json.load(fh)

    shape: Dict[str, Any] = {}
    if isinstance(data, dict):
        keys = list(data.keys())
        shape = {"type": "object", "keys": keys, "key_count": len(keys)}
    elif isinstance(data, list):
        shape = {"type": "array", "length": len(data)}
        columns: List[str] = []
        if data and all(isinstance(x, dict) for x in data):
            seen = {}
            for obj in data:
                for k in obj.keys():
                    seen[k] = True
            columns = list(seen.keys())
            shape["columns"] = columns
    else:
        shape = {"type": type(data).__name__}

    text = json.dumps(data, ensure_ascii=False, indent=2)
    meta = {"kind": "json", "encoding": enc, "shape": shape}
    return text, meta


# ---------- Public API ----------
def extract_rich(path: str) -> ExtractionResult:
    """
    Extract text + structured metadata for supported families.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(path)

    ext = _require_ext(path)
    mime = detect_mime(path)
    size = os.path.getsize(path)

    try:
        if ext in PDF_EXTS:
            text_raw, meta = _read_pdf(path)
            encoding = None
        elif ext == ".docx":
            text_raw, meta = _read_docx(path)
            encoding = None
        elif ext == ".doc":
            text_raw, meta = _read_doc(path)
            encoding = None
        elif ext in EXCEL_EXTS:
            text_raw, meta = _read_xlsx_like(path)
            encoding = None
        elif ext in CSV_EXTS:
            text_raw, meta = _read_csv_tsv_stream(path)
            encoding = meta.get("encoding")
        elif ext in TXT_EXTS:
            text_raw, meta = _read_txt(path)
            encoding = meta.get("encoding")
        elif ext in JSON_EXTS:
            text_raw, meta = _read_json(path)
            encoding = meta.get("encoding")
        else:
            raise ExtractionError(path, f"Unsupported extension '{ext}'")
    except ExtractionError:
        raise
    except Exception as e:
        logger.exception("[extract_rich] failure on %s (%s): %s", path, ext, e)
        raise ExtractionError(path, f"Failed to extract: {e}") from e

    text = _sanitize_text(text_raw, keep_newlines=True)
    stats = _basic_stats(text, size_bytes=size)

    return ExtractionResult(
        path=os.path.abspath(path),
        mime=mime,
        ext=ext,
        encoding=encoding,
        text=text,
        stats=stats,
        structure=meta,
    )


def extract(path: str, *, original_filename: Optional[str] = None) -> Dict[str, Any]:
    """
    High-level API expected by models.Document.extract_and_index():
      - returns a dict with 'content' and rich metadata (everything else)
      - includes file_name, mime, ext, encoding, size_bytes, sha256, stats, structure
    """
    res = extract_rich(path)
    try:
        sha256 = _sha256_file(path)
    except Exception:
        sha256 = ""

    meta: Dict[str, Any] = {
        "file_name": original_filename or os.path.basename(path),
        "path": res.path,
        "mime": res.mime,
        "ext": res.ext,
        "encoding": res.encoding,
        "size_bytes": res.stats.size_bytes,
        "sha256": sha256,
        "stats": {
            "num_chars": res.stats.num_chars,
            "num_words": res.stats.num_words,
            "num_lines": res.stats.num_lines,
        },
        **({"structure": res.structure} if res.structure else {}),
    }

    return {
        "content": res.text,  # LLM-ready text/TSV/pretty JSON
        **meta,               # everything else goes to Document.meta
    }


__all__ = [
    "SUPPORTED_EXTS",
    "ExtractionError",
    "ExtractionResult",
    "ExtractionStats",
    "detect_mime",
    "extract",
    "extract_rich",
]
